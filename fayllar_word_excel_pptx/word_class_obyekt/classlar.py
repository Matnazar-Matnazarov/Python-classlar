from docx import Document
class Word:
    def __init__(self, name):self.name,self.doc,self.massivv=name,Document(),[]
    def read_text(self):
        self.massivv = [i.text for i in self.doc.paragraphs]
        print('\n'.join(map(str, self.massivv)))
    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        paragraph._p = paragraph._element = None
    def write_text(self, paragraf):
        try:
            for i in range(len(self.massivv)):self.delete_paragraph(self.doc.paragraphs[0])
            self.massivv.clear()
            self.doc.add_paragraph(paragraf)
            self.massivv.append(paragraf)
        except Exception as e:  print(f'{e}')
    def add_text(self, s):
        self.doc.add_paragraph(s)
        self.massivv.append(s)
    def add_heading(self, s):
        self.doc.add_heading(s)
        self.massivv.append(s)
    def first_paragraf(self):print(self.doc.paragraphs[0].text)
    def massiv(self):print([paragraph.text for paragraph in self.doc.paragraphs])
    def del_paragraf(self, n):
        if n < len(self.doc.paragraphs): self.delete_paragraph(self.doc.paragraphs[n])
    def __repr__(self): return f"Fayl nomi: {self.name}"
    def __call__(self, n):  print(self.doc.paragraphs[n].text if n < len(self.doc.paragraphs) else "Index xato")
    def saqlash(self, s): self.doc.save(s)
    # kodni shu saytlardan o'rganib yozdim
    # https://stackoverflow.com/questions/25228106/how-to-extract-text-from-an-existing-docx-file-using-python-docx
    # https://python-docx.readthedocs.io/en/latest/
    # https://github.com/python-openxml/python-docx/tree/master/src/docx
    # https://www.geeksforgeeks.org/python-working-with-docx-module/
    # delete_paragraph bu metodni quyidagi saytdan o'rgandim 👇
    # https://stackoverflow.com/questions/29283306/removing-paragraph-from-cell-in-python-docx