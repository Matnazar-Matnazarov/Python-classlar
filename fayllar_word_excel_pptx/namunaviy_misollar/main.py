from  docx import *
# hujjat=Document() #hujjat=Document("myword.docx")
# hujjat.add_heading("salom dunyo",level=3)
# hujjat.add_paragraph("dasturchi")
# # hujjat.add_picture("usmon.jpg")
# hujjat.save("new_file.docx")
# hujjat.save("date/uz.docx")
# hujjat=Document('date/uz.docx')
# for i in hujjat.paragraphs:
#     print(i.text)
from docx import Document
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None
def funk(fayl_nomi, yangi_malumotlar):
    doc = Document(fayl_nomi)
    # doc.add_heading("salom dunyo",level=3)
    # delete_paragraph(doc.paragraphs[1])
    doc.paragraphs[0].text=yangi_malumotlar
    doc.add_paragraph("dasturcchi")
    print(len(doc.paragraphs))
    doc.save(fayl_nomi)
fayl_nomi = "new_file.docx"
yangi_malumotlar = "Men dasturchiman/qalaysan "
funk(fayl_nomi, yangi_malumotlar)
# o'rganish uchun saytlar
# https://stackoverflow.com/questions/25228106/how-to-extract-text-from-an-existing-docx-file-using-python-docx
# https://python-docx.readthedocs.io/en/latest/
# https://github.com/python-openxml/python-docx/tree/master/src/docx
# https://www.geeksforgeeks.org/python-working-with-docx-module/
